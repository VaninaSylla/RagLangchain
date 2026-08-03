from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

title = doc.add_heading("Rapport hebdomadaire de stage", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

info_lines = [
    "Stagiaire : [NOM Prénom]",
    "Établissement : [ÉCOLE / FORMATION]",
    "Entreprise d'accueil : [ENTREPRISE]",
    "Tuteur entreprise : [NOM TUTEUR]",
    "Période : semaine du 27 au 31 juillet 2026",
    "Sujet : Système RAG (Retrieval-Augmented Generation) multi-sources — Ollama + Chroma DB + Streamlit",
]
for line in info_lines:
    p = doc.add_paragraph(line)
    if line.startswith("Période") or line.startswith("Sujet"):
        p.runs[0].bold = True

doc.add_paragraph()

doc.add_heading("1. Travaux réalisés", level=1)

doc.add_heading("Lundi 27 — Mercredi 29 (Claude)", level=2)
bullets_claude = [
    ("Correction du routeur conversation", "repérage d'une faille où le classifieur basculait silencieusement vers la recherche documentaire quand l'historique était vide en début de session. Refactor pour gérer explicitement ce cas (« premier échange, rien à rappeler ») avec un generator wrapper qui unifie le streaming."),
    ("Stratégie de test du connecteur SQL", "validation en 3 niveaux (schéma seul, script automatisé posant des questions types, contrôle à l'œil de la justesse sémantique des requêtes générées) — production du fichier test_sql_connector.py."),
    ("Parser de commandes utilisateur", "module interprétant les préfixes / (/aide, /effacer, /langue, /resume, /sql, /doc, /conv, /all) et les mentions @nom_fichier pour cibler un document sans modifier le filtre sidebar."),
]
for title_b, text in bullets_claude:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(title_b + " : ")
    run.bold = True
    p.add_run(text)

doc.add_heading("Jeudi 30 — Vendredi 31 (opencode)", level=2)
bullets_opencode = [
    ("Bug embedding Ollama /tokenize", "diagnostic du port dynamique de llama-server.exe, correctif dans ingestion.py (base_url=http://localhost:11434 + timeout=120s) et OLLAMA_KEEP_ALIVE=30m côté serveur."),
    ("Déduplication des indexations par SHA-256", "nouvelles fonctions _file_sha256, purge_by_hash, list_indexed_sources ; chaque chunk reçoit metadata[source_hash] ; réindexer un PDF remplace au lieu d'empiler."),
    ("Refonte sidebar Streamlit", "source de vérité = Chroma (au lieu d'un glob du dossier), suppression de l'affichage parasite .gitkeep, stabilisation des widgets par key=."),
    ("Durcissement UI", "try/except autour de index_files avec traceback dépliable, remplacement de st.cache_resource.clear() (rerun global brutal) par load_vectorstore.clear() ciblé."),
    ("Injection sys.path dans les entry-points", "ajoutée en tête de ingest.py, rag_app.py, app_streamlit.py, cli/chat.py, cli/index.py pour permettre le lancement direct en script sans ModuleNotFoundError."),
    ("Suite de tests test_smoke.py", "~374 lignes couvrant imports, parser de commandes, validation des connecteurs SQL/Mongo, exécution SQLite in-memory, régression cp1252."),
]
for title_b, text in bullets_opencode:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(title_b + " : ")
    run.bold = True
    p.add_run(text)

doc.add_heading("2. Résultats obtenus", level=1)
p = doc.add_paragraph()
p.add_run("Pipeline RAG validé end-to-end (interrogation du PDF ").add_text("")
bold = p.add_run("LES CRYPTOMONNAIES.pdf")
bold.bold = True
p.add_run(" → 4 chunks récupérés → réponse française citée [1][2][3][4]). Base Chroma : 2 645 chunks indexés, zéro doublon après les fixes. Application Streamlit opérationnelle sur ")
code = p.add_run("http://127.0.0.1:8501")
code.font.name = "Consolas"
p.add_run(". Bilan git : 8 fichiers modifiés, +417 insertions / -47 suppressions.")

doc.add_heading("3. Difficultés rencontrées", level=1)
diff = [
    "Modèle de ports internes Ollama (llama-server.exe orphelin) difficile à diagnostiquer sans netstat combiné à l'analyse de langchain_ollama.",
    "Silent crash Streamlit (écran blanc après clic) : trois causes empilées — cache global, widgets sans key=, exception avalée par le spinner — qu'il a fallu isoler par lecture du code.",
    "__pycache__ obsolète masquant les nouveaux imports après modification de ingestion.py : résolu par nettoyage du dossier cache et relance du watcher Streamlit.",
]
for d in diff:
    doc.add_paragraph(d, style="List Bullet")

doc.add_heading("4. Objectifs de la semaine prochaine", level=1)
obj = [
    ("Coder un RAG from scratch", "implémenter sans LangChain un pipeline minimal (chunking → embeddings via Ollama → index vectoriel maison → cosine similarity → prompt + génération), pour comprendre chaque brique en profondeur avant de la ré-abstraire."),
    ("Benchmarker la version maison", "comparer la qualité de réponse et la latence contre l'implémentation LangChain actuelle."),
    ("Capitaliser sur le travail de la semaine", "intégrer test_sql_connector.py, test_smoke.py et le parser de commandes déjà produits dans la nouvelle architecture."),
    ("Identifier les briques à conserver / abandonner", "réécrire list_indexed_sources, la dédup SHA-256, l'injection sys.path si pertinents ; sinon, en tirer les leçons pour la nouvelle implémentation."),
]
for title_b, text in obj:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(title_b + " : ")
    run.bold = True
    p.add_run(text)

out = "reports/rapport_hebdomadaire_2026-07-31.docx"
doc.save(out)
print("OK ->", out)
